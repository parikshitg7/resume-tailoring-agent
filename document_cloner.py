from docx import Document

def replace_experience_bullets(resume_structure, ai_generated_experiences: list) -> Document:
    """
    Clones the original doc, replaces ONLY bullet points in the dynamic sections,
    and preserves all XML formatting.
    """
    doc = resume_structure.raw_doc
    replacement_map = {}
    
    for i, dynamic_block in enumerate(resume_structure.dynamic_blocks):
        if i >= len(ai_generated_experiences):
            break
        
        # Extract the AI data (accounting for Pydantic dict structure)
        ai_exp = ai_generated_experiences[i]
        new_bullets = ai_exp.get("bullets", [])
        new_summary = ai_exp.get("summary", "")
        
        # Find all bullet points in this specific job block
        bullet_paragraphs = [
            p for p in dynamic_block.paragraphs
            if p.text.strip().startswith(('•', '-', '–', '▪', '*')) or
               (p.paragraph_format.left_indent and p.paragraph_format.left_indent.pt > 0)
        ]
        
        # Map old bullets to new bullets
        for j, bullet_para in enumerate(bullet_paragraphs):
            if j < len(new_bullets):
                replacement_map[id(bullet_para._p)] = new_bullets[j]
        
        # Handle the 1-line summary (the first non-bullet paragraph)
        if new_summary and dynamic_block.paragraphs:
            first_para = dynamic_block.paragraphs[0]
            if not first_para.text.strip().startswith(('•', '-', '–')):
                replacement_map[id(first_para._p)] = new_summary
    
    # Perform XML Surgical Replacements
    for para in doc.paragraphs:
        para_id = id(para._p)
        if para_id in replacement_map:
            new_text = replacement_map[para_id]
            
            if para.runs:
                # Save the formatting of the very first run (font, bold, size)
                first_run = para.runs[0]
                original_format = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'font_name': first_run.font.name,
                    'font_size': first_run.font.size,
                }
                
                # Wipe all text from the XML node
                for run in para.runs:
                    run.text = ''
                
                # Inject the AI text into the first run and reapply the formatting
                first_run.text = new_text
                first_run.bold = original_format['bold']
                first_run.italic = original_format['italic']
                if original_format['font_name']:
                    first_run.font.name = original_format['font_name']
                if original_format['font_size']:
                    first_run.font.size = original_format['font_size']
                    
    return doc