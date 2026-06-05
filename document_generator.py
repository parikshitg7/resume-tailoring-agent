import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def compile_text_to_docx(markdown_text: str, output_path: str = "tailored_resume.docx"):
    """
    Parses structural markdown text from the LLM and compiles it 
    into a beautifully formatted Microsoft Word (.docx) document.
    """
    doc = docx.Document()
    
    # Configure Professional 1-Inch Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    lines = markdown_text.split("\n")
    
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue  # Skip empty whitespace lines
            
        # 1. Handle Document Main Title (H1)
        if cleaned_line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cleaned_line.replace("# ", ""))
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Arial"
            p.paragraph_format.space_after = Pt(12)
            
        # 2. Handle Major Section Headings (H2)
        elif cleaned_line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(cleaned_line.replace("## ", ""))
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Arial"
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            
        # 3. Handle Sub-headings (H3)
        elif cleaned_line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(cleaned_line.replace("### ", ""))
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Arial"
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            
        # 4. Handle Bullet Points
        elif cleaned_line.startswith("* ") or cleaned_line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Remove the markdown bullet indicator
            content = cleaned_line[2:]
            _process_inline_formatting(p, content)
            
        # 5. Handle Standard Body Paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _process_inline_formatting(p, cleaned_line)
            
    doc.save(output_path)

def _process_inline_formatting(paragraph, text: str):
    """Parses inline markdown like **bolding** and applies it to Word text runs."""
    parts = text.split("**")
    for index, part in enumerate(parts):
        run = paragraph.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        # Every odd split element means it was wrapped inside ** markers
        if index % 2 == 1:
            run.bold = True